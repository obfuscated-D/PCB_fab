from docx import Document
import re

def extract_text_from_docx(file_path, output_file):
    doc = Document(file_path)
    full_text = []
    
    # Process paragraphs with heading level and formatting
    for paragraph in doc.paragraphs:
        if not paragraph.text.strip():
            full_text.append('')  # Preserve empty paragraphs as spacing
            continue
            
        # Handle headings
        if paragraph.style.name.startswith('Heading'):
            heading_level = paragraph.style.name.split(' ')[-1]
            try:
                level = int(heading_level)
                # Create heading with appropriate # symbols
                full_text.append(f"{'#' * level} {paragraph.text}")
            except ValueError:
                # Default to formatted text if heading level isn't numeric
                full_text.append(paragraph.text)
        else:
            # Handle paragraph formatting
            text = paragraph.text
            
            # Check for bold or italic text
            has_bold = any(run.bold for run in paragraph.runs if run.bold)
            has_italic = any(run.italic for run in paragraph.runs if run.italic)
            
            # Apply basic formatting if entire paragraph has same formatting
            if has_bold and all(run.bold for run in paragraph.runs if run.text.strip()):
                text = f"**{text}**"
            if has_italic and all(run.italic for run in paragraph.runs if run.text.strip()):
                text = f"*{text}*"
                
            # Handle lists (simple detection)
            if paragraph._p.pPr and paragraph._p.pPr.numPr:
                text = f"- {text}"  # Simplistic approach for bulleted lists
                
            full_text.append(text)
    
    # Process tables with a more structured approach
    for table in doc.tables:
        # Add a separator before tables
        full_text.append("\n")
        
        # Get column widths to determine spacing
        col_count = len(table.columns)
        
        # Create a simple text-based table
        for i, row in enumerate(table.rows):
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.replace('\n', ' ')  # Remove internal newlines
                row_text.append(cell_text)
            
            # Create a simple pipe-separated format for the table
            table_row = " | ".join(row_text)
            full_text.append(f"| {table_row} |")
            
            # Add separator line after header row
            if i == 0:
                separator = "|" + "|".join([" --- " for _ in range(col_count)]) + "|"
                full_text.append(separator)
        
        # Add spacing after table
        full_text.append("\n")
    
    # Write to output file
    with open(output_file, "w", encoding="utf-8") as text_file:
        text_file.write('\n'.join(full_text))
    
    print(f"Conversion complete. Text saved to {output_file}")


if __name__ == "__main__":
    import sys
    if len(sys.argv) != 3:
        print("Usage: python docx_to_txt.py <input.docx> <output.txt>")
    else:
        input_file = sys.argv[1]
        output_file = sys.argv[2]
        extract_text_from_docx(input_file, output_file)